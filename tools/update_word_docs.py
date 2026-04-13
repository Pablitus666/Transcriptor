import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def update_docx(file_path):
    if not os.path.exists(file_path):
        print(f"Skipping: {file_path}")
        return
    
    try:
        doc = Document(file_path)
        
        # Añadir un salto de página y una sección de instalación modular al final
        doc.add_page_break()
        
        title = doc.add_paragraph("🚀 INSTRUCCIONES DE INSTALACIÓN MODULAR")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.runs[0]
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(2, 48, 71) # Azul oscuro profesional

        instr = doc.add_paragraph()
        instr.add_run("\nPara garantizar una instalación eficiente de los componentes de Inteligencia Artificial (IA), siga estos pasos obligatorios:\n\n").bold = True
        
        p1 = doc.add_paragraph("1. ")
        p1.add_run("Instalación Base: ").bold = True
        p1.add_run("Ejecute 'Transcriptor_Setup.exe' para instalar la estructura del programa en C:\\Transcriptor.")
        
        p2 = doc.add_paragraph("2. ")
        p2.add_run("Copia del Motor (CRÍTICO): ").bold = True
        p2.add_run("Copie las carpetas 'whisper_env' y 'models_cache' desde su disco de distribución a la carpeta de instalación (C:\\Transcriptor).")
        
        p3 = doc.add_paragraph("3. ")
        p3.add_run("Inicio: ").bold = True
        p3.add_run("Use el acceso directo del escritorio para iniciar el sistema.")
        
        # Guardar cambios
        doc.save(file_path)
        print(f"✅ Documento actualizado: {file_path}")
    except Exception as e:
        print(f"❌ Error al actualizar {file_path}: {e}")

docs = [
    "CAPACITACION_TECNICA.docx",
    "MANUAL.docx",
    "📘 GUÍA INSTITUCIONAL.docx",
    "GUIA_INSTITUCIONAL.docx"
]

for d in docs:
    update_docx(d)
