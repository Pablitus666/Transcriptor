import os
from docx import Document

def fix_docx_name(file_path):
    if not os.path.exists(file_path):
        return
    
    try:
        doc = Document(file_path)
        replaced = False
        
        # 1. Buscar y reemplazar en párrafos
        for p in doc.paragraphs:
            if "Transcriptor_Setup_Base" in p.text:
                p.text = p.text.replace("Transcriptor_Setup_Base", "Transcriptor_Setup")
                replaced = True
        
        # 2. Buscar y reemplazar en tablas (por si acaso)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "Transcriptor_Setup_Base" in p.text:
                            p.text = p.text.replace("Transcriptor_Setup_Base", "Transcriptor_Setup")
                            replaced = True
        
        if replaced:
            doc.save(file_path)
            print(f"✅ Documento corregido: {file_path}")
        else:
            print(f"ℹ️ Sin cambios en: {file_path}")
            
    except Exception as e:
        print(f"❌ Error en {file_path}: {e}")

docs = [
    "CAPACITACION_TECNICA.docx",
    "MANUAL.docx",
    "📘 GUÍA INSTITUCIONAL.docx",
    "GUIA_INSTITUCIONAL.docx"
]

for d in docs:
    fix_docx_name(d)
