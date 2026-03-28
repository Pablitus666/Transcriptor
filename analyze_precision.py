from docx import Document
import difflib
import re

def normalize_text(text):
    # Eliminar espacios múltiples y normalizar caracteres invisibles
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def extract_content(file_path):
    doc = Document(file_path)
    # Extraer texto de párrafos y tablas (las tablas suelen contener los diálogos)
    content = []
    
    # Primero párrafos normales
    for p in doc.paragraphs:
        if p.text.strip():
            content.append(normalize_text(p.text))
            
    # Luego tablas (donde suelen estar los diálogos de Cámara Gesell)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([normalize_text(cell.text) for cell in row.cells])
            if row_text.strip():
                content.append(row_text)
                
    return content

def compare_transcripts(script_file, corrected_file):
    print(f"Leyendo: {script_file}")
    script_content = extract_content(script_file)
    
    print(f"Leyendo: {corrected_file}")
    corrected_content = extract_content(corrected_file)
    
    # Unión de todo el contenido para una comparación global
    script_full = "\n".join(script_content)
    corrected_full = "\n".join(corrected_content)
    
    # Calcular similitud básica
    similarity = difflib.SequenceMatcher(None, script_full, corrected_full).ratio()
    print(f"\nNivel de Similitud Global: {similarity*100:.2f}%\n")
    
    # Mostrar las primeras 5 diferencias significativas
    print("Principales diferencias detectadas (Script vs Corregido):")
    diff = list(difflib.ndiff(script_content[:20], corrected_content[:20]))
    
    count = 0
    for line in diff:
        if (line.startswith('- ') or line.startswith('+ ')) and count < 10:
            print(line)
            count += 1

if __name__ == "__main__":
    file_script = r"output/Sandra_transcrito.docx"
    file_corrected = r"output/Sandra transcrito corregido.docx"
    compare_transcripts(file_script, file_corrected)
