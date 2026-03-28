from docx import Document
import os

def fix_guia_final():
    path = "📘 GUÍA INSTITUCIONAL.docx"
    if not os.path.exists(path): return
    doc = Document(path)
    
    # Buscamos y reemplazamos los párrafos específicos de la sección 3
    for p in doc.paragraphs:
        # Detectamos la línea de "dos formatos"
        if "exportar en dos formatos" in p.text:
            p.text = p.text.replace("exportar en dos formatos:", "exportar en formato profesional:")
        
        # Eliminamos la línea del TXT
        if "TXT: Una transcripción" in p.text:
            p.text = "" # La dejamos vacía para luego limpiar
            
        # Mejoramos la descripción del DOCX
        if "DOCX: Un documento" in p.text:
            p.text = "• DOCX: Un documento formateado institucionalmente (Arial 11) que permite el uso de plantillas personalizadas."

    # Limpieza de párrafos vacíos resultantes de la eliminación del TXT
    paragraphs_to_keep = [p.text for p in doc.paragraphs if p.text.strip()]
    
    # Re-creamos el documento con el contenido limpio (esto es más seguro para párrafos pequeños)
    # Pero como queremos mantener el formato del resto, mejor operamos sobre los elementos oxml si es necesario.
    # Intentemos simplemente borrar el elemento del párrafo vacío.
    
    for p in list(doc.paragraphs):
        if p.text == "":
            p._element.getparent().remove(p._element)

    doc.save(path)
    print(f"✅ {path} corregido definitivamente.")

if __name__ == "__main__":
    fix_guia_final()
