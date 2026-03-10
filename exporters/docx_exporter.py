# exporters/docx_exporter.py
import re
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- MOTOR DE BÚSQUEDA PROFESIONAL (Basado en el motor de Wordy) ---
class RunMatcher:
    """
    Especialista en mapeo y manipulación quirúrgica de Runs en DOCX.
    Normaliza el texto para asegurar coincidencias con espacios de Word (\xa0).
    """
    def __init__(self, paragraph, pattern):
        self.paragraph = paragraph
        self.pattern = pattern
        self.run_map = []
        self.full_text = ""
        self._build_map()

    def _build_map(self):
        """Construye el mapa de texto normalizando espacios de Word."""
        pos = 0
        for run in self.paragraph.runs:
            text = run.text if run.text else ""
            # Sincronización con Wordy: manejo de \xa0
            normalized_text = text.replace('\xa0', ' ')
            length = len(normalized_text)
            if length == 0: continue
            self.run_map.append({
                "run": run, 
                "start": pos, 
                "end": pos + length
            })
            self.full_text += normalized_text
            pos += length

    def format_matches(self):
        """Busca coincidencias y aplica negrita + subrayado."""
        if not self.full_text: return
        matches = list(self.pattern.finditer(self.full_text))
        # Aplicamos en reversa para que los splits no afecten los offsets
        for match in reversed(matches):
            start, end = match.span()
            self._apply_surgical_format(start, end)

    def _apply_surgical_format(self, start, end):
        affected = [e for e in self.run_map if e["end"] > start and e["start"] < end]
        if not affected: return

        # Split al inicio (quirúrgico)
        first = affected[0]
        if start > first["start"]:
            new_run = self._split_run(first["run"], start - first["start"])
            first["run"] = new_run
            first["start"] = start

        # Split al final (quirúrgico)
        last = affected[-1]
        if end < last["end"]:
            self._split_run(last["run"], end - last["start"])
            last["end"] = end

        # Aplicar Negrita y Subrayado
        for entry in affected:
            run = entry["run"]
            run.bold = True
            run.underline = True

    def _split_run(self, run, split_point):
        """Divide un run preservando el estilo exacto (Arial 11)."""
        text = run.text
        run.text = text[:split_point]
        new_run = self.paragraph.add_run(text[split_point:])
        # Herencia de estilo profesional
        new_run.font.name = run.font.name
        new_run.font.size = run.font.size
        new_run.font.bold = run.font.bold
        new_run.font.italic = run.font.italic
        new_run.font.underline = run.font.underline
        if run.font.color and run.font.color.rgb:
            new_run.font.color.rgb = run.font.color.rgb
        
        run._r.addnext(new_run._r)
        return new_run

def _compile_super_regex(words_list):
    """
    Compila una regex ultra-robusta que EXIGE los dos puntos (:) al final.
    Ignora palabras que no tengan el signo ':'.
    """
    vowel_map = {
        'a': '[aá]', 'á': '[aá]',
        'e': '[eé]', 'é': '[eé]',
        'i': '[ií]', 'í': '[ií]',
        'o': '[oó]', 'ó': '[oó]',
        'u': '[uúü]', 'ú': '[uúü]'
    }

    escaped_parts = []
    for w in words_list:
        # Quitamos cualquier ':' que venga en la lista para procesar la raíz
        w_clean = w.replace(':', '')
        escaped_w = re.escape(w_clean)
        
        gender_match = re.search(r'([oaóá])$', escaped_w, flags=re.IGNORECASE)
        gender_pos = gender_match.start(1) if gender_match else -1
        
        final_w = ""
        for i, char in enumerate(escaped_w):
            lower_char = char.lower()
            if i == gender_pos:
                final_w += '[aáoó]'
            elif lower_char in vowel_map:
                final_w += vowel_map[lower_char]
            else:
                final_w += char
        escaped_parts.append(final_w)
    
    pattern_str = '|'.join(escaped_parts)
    
    # EXIGENCIA DE DOS PUNTOS (:): 
    # La regex busca (Palabra)(:) de forma obligatoria.
    # El boundary permite detectar palabras al puro inicio del párrafo.
    boundary = r'(?<![a-zA-Z0-9áéíóúÁÉÍÓÚñÑ])'
    
    # Pattern: (Palabra con variantes de tildes/género) seguido inmediatamente de ':'
    return re.compile(rf'{boundary}(({pattern_str}):)', re.IGNORECASE)

def add_page_number(paragraph):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    paragraph.add_run('')._r.append(fldChar1)
    paragraph.add_run('')._r.append(instrText)
    paragraph.add_run('')._r.append(fldChar2)

def export_to_docx(segments, output_path, template_path=None):
    """
    Genera un archivo DOCX con inteligencia lingüística enfocada.
    Resalta hablantes en párrafos, encabezados y pies de página, pero ignora TABLAS.
    """
    # Lista de palabras clave base (el motor expande tildes, géneros y exige ':')
    keywords = ["Psicologo", "Psicologa", "Victima"]
    bold_pattern = _compile_super_regex(keywords)

    if template_path:
        doc = Document(template_path)
        placeholder_found = False
        
        # 1. Inserción de la transcripción en el marcador {{TRANSCRIPCION}}
        for p in list(doc.paragraphs):
            if "{{TRANSCRIPCION}}" in p.text:
                style = p.style
                alignment = p.alignment
                
                for i, segment in enumerate(segments):
                    text_content = f"{segment['speaker']}: {segment['text']}"
                    text_p = p.insert_paragraph_before(text_content, style)
                    text_p.alignment = alignment

                    # Forzar estilo Arial 11 en la inserción
                    for run in text_p.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                    
                    if i < len(segments) - 1:
                        empty_p = p.insert_paragraph_before("", style)
                        run = empty_p.add_run(' ')
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)

                p._element.getparent().remove(p._element)
                placeholder_found = True
        
        # 2. ESCANEO QUIRÚRGICO (Ignora Tablas):
        # Procesamos solo los párrafos libres del documento, encabezados y pies.
        
        # A. Cuerpo (Donde está la transcripción recién insertada)
        for p in doc.paragraphs:
            matcher = RunMatcher(p, bold_pattern)
            matcher.format_matches()

        # B. Secciones (Encabezados y Pies de página - SOLO PÁRRAFOS LIBRES)
        for section in doc.sections:
            # Procesar párrafos del encabezado (pero no sus tablas)
            for p in section.header.paragraphs:
                matcher = RunMatcher(p, bold_pattern)
                matcher.format_matches()
                
            # Procesar párrafos del pie de página (pero no sus tablas)
            for p in section.footer.paragraphs:
                matcher = RunMatcher(p, bold_pattern)
                matcher.format_matches()
        
        if not placeholder_found:
            for s in segments:
                p = doc.add_paragraph(f"{s['speaker']}: {s['text']}")
                matcher = RunMatcher(p, bold_pattern)
                matcher.format_matches()

    else:
        doc = Document()
        # ... (Configuración de página por defecto) ...
        section = doc.sections[0]
        section.page_height = Mm(330); section.page_width = Mm(216)
        section.left_margin = Cm(2); section.right_margin = Cm(2)
        section.top_margin = Cm(2); section.bottom_margin = Cm(2)
        
        p_footer = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_number(p_footer)

        doc.add_heading('Transcripción', level=1)
        doc.add_paragraph()
        
        for s in segments:
            par = doc.add_paragraph()
            run = par.add_run(f"{s['speaker']}: {s['text']}")
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Aplicar motor de resaltado
            matcher = RunMatcher(par, bold_pattern)
            matcher.format_matches()

    doc.save(output_path)
