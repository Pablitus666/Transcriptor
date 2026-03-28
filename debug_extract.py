from docx import Document
import re

def normalize_text(text):
    return re.sub(r'\s+', ' ', text).strip()

def extract_to_txt(docx_path, txt_path):
    doc = Document(docx_path)
    with open(txt_path, 'w', encoding='utf-8') as f:
        # Párrafos
        for p in doc.paragraphs:
            if p.text.strip():
                f.write(normalize_text(p.text) + "\n")
        # Tablas
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([normalize_text(cell.text) for cell in row.cells])
                if row_text.strip():
                    f.write(row_text + "\n")

if __name__ == "__main__":
    extract_to_txt(r"output/Sandra_transcrito.docx", "script_text.txt")
    extract_to_txt(r"output/Sandra transcrito corregido.docx", "corrected_text.txt")
    print("Archivos extraídos: script_text.txt y corrected_text.txt")
