from docx import Document
import difflib
import re

def normalize_text(text):
    return re.sub(r'\s+', ' ', text).strip()

def get_stats(file_path):
    doc = Document(file_path)
    paragraphs = len(doc.paragraphs)
    tables = len(doc.tables)
    total_cells = sum(len(row.cells) for table in doc.tables for row in table.rows)
    return f"Párrafos: {paragraphs}, Tablas: {tables}, Celdas: {total_cells}"

def extract_raw_text(file_path):
    doc = Document(file_path)
    text = []
    # Extraer de párrafos
    for p in doc.paragraphs:
        if p.text.strip():
            text.append(normalize_text(p.text))
    # Extraer de tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text.append(normalize_text(cell.text))
    return text

def compare_transcripts(script_file, corrected_file):
    print(f"Stats Script: {get_stats(script_file)}")
    print(f"Stats Corrected: {get_stats(corrected_file)}")
    
    script_content = extract_raw_text(script_file)
    corrected_content = extract_raw_text(corrected_file)
    
    script_full = " ".join(script_content)
    corrected_full = " ".join(corrected_content)
    
    similarity = difflib.SequenceMatcher(None, script_full, corrected_full).ratio()
    print(f"\nNivel de Similitud Global (Texto bruto): {similarity*100:.2f}%\n")
    
    # Mostrar una muestra del texto de ambos
    print("--- MUESTRA SCRIPT ---")
    print("\n".join(script_content[:10]))
    print("\n--- MUESTRA CORREGIDO ---")
    print("\n".join(corrected_content[:10]))

if __name__ == "__main__":
    file_script = r"output/Sandra_transcrito.docx"
    file_corrected = r"output/Sandra transcrito corregido.docx"
    compare_transcripts(file_script, file_corrected)
